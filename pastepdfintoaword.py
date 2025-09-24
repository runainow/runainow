import os
import fitz  # PyMuPDF
import docx
from docx.shared import Inches
# it works on windows, do not forget pip install PyMuPDF
# pip install python-docx
# --- 1. Configuration ---
INPUT_FOLDER = 'in'
WORD_TEMPLATE = 'doc1a.docx'
OUTPUT_WORD_FILE = 'doc1a_completed_all_final.docx'
PLACEHOLDER_TEXT = '黏貼證照影本'
TEMP_IMAGE_FILE = 'temp_certificate_image.png'

def main():
    """Main function to insert ALL PDFs, auto-extending the table as needed."""
    
    # --- 2. Find and sort all PDF files ---
    print(f"Scanning for PDFs in './{INPUT_FOLDER}/' folder...")
    if not os.path.isdir(INPUT_FOLDER):
        print(f"❌ Error: Input folder '{INPUT_FOLDER}' not found. Please create it.")
        return

    pdfs_to_process = sorted([f for f in os.listdir(INPUT_FOLDER) if f.lower().endswith('.pdf')])
    
    if not pdfs_to_process:
        print(f"❌ No PDF files found in the '{INPUT_FOLDER}' folder.")
        return
    
    print(f"✅ Found {len(pdfs_to_process)} PDF(s) to process.")

    # --- 3. Open Word doc and prepare table cells ---
    try:
        doc = docx.Document(WORD_TEMPLATE)
        if not doc.tables:
            print(f"❌ Error: No tables found in '{WORD_TEMPLATE}'.")
            return
        
        main_table = max(doc.tables, key=lambda t: len(t._cells))
        print(f"✅ Target table identified with {len(main_table.rows)} rows.")
        
        # Collect existing cells from every second row (2nd, 4th, 6th...)
        cells_to_populate = []
        for i, row in enumerate(main_table.rows):
            if i % 2 != 0: # Target 2nd, 4th, 6th rows (index 1, 3, 5...)
                cells_to_populate.extend(row.cells)
        
        # --- NEW: Auto-extend the table if more space is needed ---
        num_pdfs = len(pdfs_to_process)
        num_available_cells = len(cells_to_populate)
        
        if num_pdfs > num_available_cells:
            num_columns = len(main_table.columns) if main_table.rows else 2
            cells_needed = num_pdfs - num_available_cells
            # Calculate how many new image rows we need (rounding up)
            rows_to_add = (cells_needed + num_columns - 1) // num_columns
            
            print(f"⚠️ Not enough space. Adding {rows_to_add * 2} new rows to the table...")
            
            for _ in range(rows_to_add):
                # Add a new row for placeholder text
                text_row = main_table.add_row()
                for i in range(num_columns):
                    text_row.cells[i].text = PLACEHOLDER_TEXT
                
                # Add a new row for images and add its cells to our target list
                image_row = main_table.add_row()
                cells_to_populate.extend(image_row.cells)

        print(f"✅ Table is ready with {len(cells_to_populate)} total slots for images.")

    except Exception as e:
        print(f"❌ Error preparing Word document: {e}")
        return

    # --- 4. Process each PDF and insert into a target cell ---
    num_inserted = 0
    for pdf_filename, cell in zip(pdfs_to_process, cells_to_populate):
        pdf_path = os.path.join(INPUT_FOLDER, pdf_filename)
        print(f"  -> Processing '{pdf_filename}'...")
        try:
            with fitz.open(pdf_path) as pdf_doc:
                page = pdf_doc.load_page(0)
                pix = page.get_pixmap(dpi=300)
                pix.save(TEMP_IMAGE_FILE)

            cell.text = ''
            run = cell.paragraphs[0].add_run()
            run.add_picture(TEMP_IMAGE_FILE, width=Inches(3.0))
            num_inserted += 1
        except Exception as e:
            print(f"  ⚠️ Could not process '{pdf_filename}': {e}")
        finally:
            if os.path.exists(TEMP_IMAGE_FILE):
                os.remove(TEMP_IMAGE_FILE)
    
    # --- 5. Save the final document ---
    if num_inserted > 0:
        doc.save(OUTPUT_WORD_FILE)
        print(f"\n🎉 Successfully inserted {num_inserted} certificate(s).")
        print(f"✅ Final document saved as '{OUTPUT_WORD_FILE}'.")
    else:
        print("\nNo certificates were inserted.")

if __name__ == "__main__":
    main()